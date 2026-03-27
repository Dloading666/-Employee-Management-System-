from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "企业员工考察管理系统技术讲解文档.docx"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_run(run, font_name="宋体", size=Pt(10.5), bold=False, color=None) -> None:
    set_east_asia_font(run, font_name)
    run.font.size = size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, first_line_chars: int | None = None, after=Pt(6)) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = after
    fmt.space_before = Pt(0)
    if first_line_chars:
        fmt.first_line_indent = Pt(first_line_chars * 10.5)


def add_text_paragraph(document, text: str, *, first_line_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = document.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    style_run(run)
    style_paragraph(p, first_line_chars=first_line_chars)
    return p


def add_bullet(document, text: str):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    style_run(run)
    style_paragraph(p, first_line_chars=None, after=Pt(3))
    return p


def add_number(document, text: str):
    p = document.add_paragraph(style="List Number")
    run = p.add_run(text)
    style_run(run)
    style_paragraph(p, first_line_chars=None, after=Pt(3))
    return p


def set_cell_text(cell, text: str, *, bold=False, font_name="宋体", size=Pt(10.5), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    style_paragraph(paragraph, first_line_chars=None, after=Pt(0))
    run = paragraph.add_run(text)
    style_run(run, font_name=font_name, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = tbl_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "B4C7DC")


def add_table(document, headers, rows, *, column_widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(
            cell,
            header,
            bold=True,
            font_name="微软雅黑",
            size=Pt(10.5),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(cell, "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    if column_widths:
        for row in table.rows:
            for index, width_cm in enumerate(column_widths):
                row.cells[index].width = Cm(width_cm)
    document.add_paragraph()
    return table


def apply_global_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, size, color in [
        ("Title", Pt(24), RGBColor(33, 79, 121)),
        ("Heading 1", Pt(16), RGBColor(31, 78, 121)),
        ("Heading 2", Pt(13), RGBColor(31, 78, 121)),
        ("Heading 3", Pt(12), RGBColor(68, 114, 196)),
    ]:
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = color


def add_heading(document, text: str, level: int) -> None:
    p = document.add_heading("", level=level)
    run = p.add_run(text)
    style_run(
        run,
        font_name="微软雅黑",
        size={1: Pt(16), 2: Pt(13), 3: Pt(12)}.get(level, Pt(11)),
        bold=True,
        color={1: RGBColor(31, 78, 121), 2: RGBColor(31, 78, 121), 3: RGBColor(68, 114, 196)}.get(level),
    )


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("企业员工考察管理系统技术讲解文档")
    style_run(run, font_name="微软雅黑", size=Pt(24), bold=True, color=RGBColor(31, 78, 121))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于当前源码的系统设计、业务流程与实现说明")
    style_run(run, font_name="微软雅黑", size=Pt(13), color=RGBColor(89, 89, 89))

    meta_lines = [
        "项目代号：attendance-management / frontend-vue",
        f"生成日期：{date.today().isoformat()}",
        "适用场景：课程答辩、项目汇报、技术交接、二次开发说明",
        "说明：本文档根据当前工作区中的 Spring Boot + Vue 源码整理，不依赖空泛模板。",
    ]
    for line in meta_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        style_run(run, font_name="宋体", size=Pt(11))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("企业员工考察管理系统当前处于前后端分离版本，后端采用 Session 鉴权，前端使用 Vue 3 单页应用。")
    style_run(run, size=Pt(11), color=RGBColor(80, 80, 80))

    document.add_page_break()


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("企业员工考察管理系统技术讲解文档")
    style_run(run, font_name="微软雅黑", size=Pt(9), color=RGBColor(127, 127, 127))


def build_document() -> Document:
    document = Document()
    apply_global_styles(document)
    add_footer(document.sections[0])

    document.core_properties.title = "企业员工考察管理系统技术讲解文档"
    document.core_properties.author = "Codex"
    document.core_properties.subject = "Java Web 项目技术讲解"
    document.core_properties.comments = "基于当前工作区源码自动整理生成"

    add_cover(document)

    add_heading(document, "1. 项目概述与文档目标", 1)
    add_text_paragraph(
        document,
        "企业员工考察管理系统是一套面向企业基础人事管理场景的业务系统，主要覆盖员工信息、部门岗位、考勤打卡、请假审批、薪资记录和个人资料维护等功能。当前代码已经完成从传统 SSM/XML 风格向 Spring Boot 3 + Vue 3 前后端分离模式的迁移，系统整体具备比较清晰的分层结构，适合作为课程设计项目、毕业设计展示项目以及小型企业信息化原型系统。",
    )
    add_text_paragraph(
        document,
        "本文档的目标不是简单罗列技术名词，而是结合源码实际实现，说明该系统为什么这样设计、模块之间如何协作、数据与权限如何流转，以及当前版本的优势、边界和后续可优化方向。阅读完本文后，读者应当能够独立理解该项目的整体架构，并对后续维护和扩展具备明确抓手。",
    )
    for item in [
        "从业务角度解释系统解决了什么问题。",
        "从技术角度解释前端、后端、数据库和安全组件如何配合。",
        "从工程角度解释当前实现的亮点、局限与演进方向。",
    ]:
        add_bullet(document, item)

    add_heading(document, "2. 总体架构说明", 1)
    add_text_paragraph(
        document,
        "系统采用典型的前后端分离架构。浏览器访问 Vue 单页应用，页面通过 Axios 调用 `/api/**` 下的 REST 接口；后端 Spring Boot 负责请求分发、会话鉴权、业务逻辑处理和数据库访问；持久层以 MyBatis Mapper + XML 为主，统计和薪资等少量场景直接使用 JdbcTemplate 进行 SQL 查询。",
    )
    add_text_paragraph(
        document,
        "简化后的调用链路如下：浏览器页面 -> Vue Router -> Axios -> Spring Security + SessionAuthenticationFilter -> Controller -> Service -> Mapper / JdbcTemplate -> MySQL -> JSON 响应 -> 前端页面渲染。",
        first_line_chars=0,
    )
    add_table(
        document,
        ["层次", "主要技术", "职责", "代表位置"],
        [
            ["表现层", "Vue 3、Vue Router、Bootstrap、ECharts", "渲染页面、路由控制、表单交互、统计图表展示", "frontend-vue/src/pages、frontend-vue/src/layouts"],
            ["接口层", "Spring MVC、REST Controller", "接收请求、参数校验、封装统一响应、控制访问边界", "src/main/java/edu/com/eams/controller/api"],
            ["安全层", "Spring Security、SessionAuthenticationFilter", "基于 Session 恢复登录态、区分管理员与员工权限", "src/main/java/edu/com/eams/config、security"],
            ["业务层", "Service", "封装员工、考勤、请假、用户等核心业务逻辑", "src/main/java/edu/com/eams/service"],
            ["数据访问层", "MyBatis Mapper/XML、JdbcTemplate", "执行 CRUD、分页查询、统计 SQL 和部分直接写库操作", "src/main/resources/mapper、SalaryApiController、DashboardApiController"],
            ["存储层", "MySQL 8", "保存用户、员工、部门、岗位、考勤、请假、薪资等业务数据", "application.yml 中的 DB_URL 配置"],
        ],
        column_widths=[2.6, 4.6, 6.2, 3.2],
    )

    add_heading(document, "3. 技术栈与选型分析", 1)
    add_table(
        document,
        ["类别", "当前选型", "作用", "选型评价"],
        [
            ["后端框架", "Spring Boot 3.2.6", "提供自动装配、Web 容器、配置管理和工程化启动方式", "替代旧式 `web.xml` 启动，降低部署和维护成本"],
            ["安全框架", "Spring Security", "保护 `/api/**` 接口、处理认证失败和权限不足场景", "适合与 Session 结合，实施成本低"],
            ["持久层", "MyBatis 3 + XML Mapper", "完成实体映射、分页列表、条件查询与 CRUD", "保留传统 SQL 可控性，迁移成本较低"],
            ["数据库连接池", "Druid", "管理数据库连接、提升并发访问稳定性", "适合教学项目和中小规模业务系统"],
            ["前端框架", "Vue 3 + Vite", "构建单页应用并提升开发时热更新效率", "开发体验好，适合逐步替换旧页面"],
            ["前端通信", "Axios", "发起 REST 请求并统一处理响应结构", "与 `RetJson` 结构配合简单直接"],
            ["UI/图表", "Bootstrap 5 + ECharts 6", "完成后台管理界面与数据可视化", "上手快，适合管理系统场景"],
            ["构建工具", "Maven、Vite", "后端 WAR 打包，前端产出静态资源", "职责清晰，便于分开部署与调试"],
        ],
        column_widths=[2.3, 4.2, 4.8, 5.3],
    )
    add_text_paragraph(
        document,
        "从当前实现来看，项目技术选型强调“低门槛、可迁移、便于展示”。后端并未引入复杂的微服务、中间件或消息队列，而是围绕单体应用做了合理分层；前端也优先采用成熟且学习成本较低的 Vue + Bootstrap 组合。这使项目非常适合用于教学汇报、项目答辩和后续二次开发。",
    )

    add_heading(document, "4. 功能模块讲解", 1)
    add_heading(document, "4.1 管理员侧模块", 2)
    add_table(
        document,
        ["模块", "典型路由", "核心接口", "功能说明"],
        [
            ["首页统计", "/app/welcome", "/api/dashboard/**", "展示用户数、员工数、部门数、今日考勤，以及部门占比、请假与新增趋势图"],
            ["部门管理", "/app/departments", "/api/departments", "维护部门名称、负责人、层级和状态"],
            ["岗位管理", "/app/postlevels", "/api/postlevels", "维护岗位等级、基础薪资、排序信息"],
            ["员工管理", "/app/employees", "/api/employees", "管理员工基础资料、入职信息、岗位和部门归属"],
            ["用户管理", "/app/users", "/api/users", "维护系统登录账号、状态和密码配置"],
            ["考勤管理", "/app/attendance", "/api/attendance", "管理员可查看和维护所有员工考勤记录"],
            ["薪资管理", "/app/salary-admin", "/api/salaries", "按月份维护员工薪资，支持管理员录入和删除"],
            ["请假审批", "/app/leave-approval", "/api/leaves、/api/leaves/{id}/status", "查看员工请假单并进行单条或批量审批"],
        ],
        column_widths=[2.6, 3.4, 4.6, 6.0],
    )
    add_heading(document, "4.2 员工侧模块", 2)
    add_table(
        document,
        ["模块", "典型路由", "核心接口", "功能说明"],
        [
            ["登录与退出", "/signin", "/api/auth/login、/api/auth/logout、/api/auth/me", "通过 Session 建立和恢复登录态"],
            ["个人主页", "/app/welcome", "/api/dashboard/**", "员工登录后也可看到首页统计概览"],
            ["个人信息", "/app/profile", "/api/profile", "查看和维护与本人绑定的员工档案"],
            ["我的考勤", "/app/attendance", "/api/attendance、/api/attendance/check-in、/check-out", "执行上班打卡、下班签退并查看本人历史记录"],
            ["请假申请", "/app/leave-apply", "/api/leaves", "提交请假时间、请假类型和请假原因"],
            ["请假记录", "/app/leave-list", "/api/leaves", "查看本人请假单及审批结果"],
            ["我的薪资", "/app/salary-my", "/api/salaries/mine", "按月份查看个人薪资记录"],
        ],
        column_widths=[2.6, 3.4, 4.8, 5.8],
    )
    add_text_paragraph(
        document,
        "从菜单与路由的划分可以看出，系统采用“同一套前端壳体 + 角色分流菜单”的思路。管理员与员工共用同一个 AppShell，但通过 Vue Router 的 `meta.adminOnly` 和后端的角色校验共同控制可访问范围，这种做法实现简单且足够清晰。",
    )

    add_heading(document, "5. 核心数据模型", 1)
    add_text_paragraph(
        document,
        "由于工作区中未提供初始化 SQL 文件，以下数据模型说明主要依据实体类、Mapper XML 和 Controller 中的查询语句整理而成，因此足以支撑架构理解，但不等同于最终数据库设计说明书。",
    )
    add_table(
        document,
        ["表 / 实体", "核心字段示例", "业务作用"],
        [
            ["sys_user / User", "id、name、account、password、status", "系统登录账号表，决定用户是否能进入系统"],
            ["emp / Employee", "name、code、email、mobile、entry_date、dept_id、post_level_id、user_id", "员工主档案，记录人员详细信息并与系统账号关联"],
            ["sys_dept / Department", "name、manager_user_id、parent_id、sort_code、status", "部门组织结构与负责人信息"],
            ["emp_post_level / PostLevel", "name、level、salary、sort_code、status", "岗位等级与基础薪资配置"],
            ["emp_attendance_record / Attendance", "user_id、attendance_date、check_in_time、check_out_time、status", "每日考勤记录，保存签到签退时间"],
            ["emp_leave_record / LeaveRecord", "user_id、leave_start_time、leave_end_time、leave_reason、approve_user_id、status", "请假申请与审批流转记录"],
            ["emp_salary", "user_id、user_name、pay_month、amount、remark、create_time", "按月份保存员工薪资条目"],
        ],
        column_widths=[4.2, 6.0, 6.4],
    )
    add_text_paragraph(
        document,
        "其中，`sys_user` 与 `emp` 的关系是系统理解的关键。系统登录依赖 `sys_user`，而员工业务档案依赖 `emp`；`emp.user_id` 负责把“登录身份”与“员工信息”绑定起来，因此个人资料、我的考勤、我的薪资等能力才能按照当前会话准确过滤出本人数据。",
    )

    add_heading(document, "6. 认证与权限控制机制", 1)
    add_heading(document, "6.1 登录过程", 2)
    add_number(document, "前端登录页将账号和密码提交到 `/api/auth/login`。")
    add_number(document, "后端 `LoginServiceImpl` 根据账号查询 `sys_user`，再通过 `PasswordSupport` 判断密码是否匹配。")
    add_number(document, "若密码是 BCrypt 形式，则使用 `PasswordEncoder` 校验；若仍是明文，则直接比较字符串。这意味着系统支持从旧密码存储方式向 BCrypt 逐步迁移。")
    add_number(document, "登录成功后，后端会对用户对象做脱敏处理，只保留必要字段，并把 `currentUser` 与 `isAdmin` 写入 Session。")
    add_heading(document, "6.2 请求鉴权", 2)
    add_text_paragraph(
        document,
        "`SessionAuthenticationFilter` 会在每次请求时读取 Session 中的 `currentUser`，并补充 Spring Security 所需的认证对象。管理员会被授予 `ROLE_ADMIN` 和 `ROLE_EMPLOYEE` 两个角色，普通员工只拥有 `ROLE_EMPLOYEE`。之后由 `SecurityConfig` 决定不同接口的访问策略。",
    )
    add_text_paragraph(
        document,
        "例如，`/api/employees`、`/api/departments`、`/api/postlevels`、`/api/users`、`/api/salaries` 仅允许管理员访问；`/api/profile/**`、`/api/leaves/**`、`/api/attendance/**` 则要求用户已登录。对前端来说，只要浏览器携带 Session Cookie，就能被后端正确识别。",
    )
    add_heading(document, "6.3 当前权限方案的特点", 2)
    for item in [
        "优点：实现成本低，适合单体项目和课程设计项目快速落地。",
        "优点：无需额外发放 JWT，前后端通过 `withCredentials` 即可维持会话。",
        "边界：当前管理员身份本质上仍依赖 `account = admin` 的兼容判断，并非真正的 RBAC 模型。",
        "建议：后续可扩展为 `用户-角色-权限` 三张表，提升可维护性和可扩展性。",
    ]:
        add_bullet(document, item)

    add_heading(document, "7. 关键业务流程解析", 1)
    add_heading(document, "7.1 考勤打卡流程", 2)
    add_text_paragraph(
        document,
        "员工点击“上班打卡”时，前端调用 `/api/attendance/check-in`。服务层先按“当前用户 + 今天日期”查询是否已有考勤记录：若没有，则插入一条新记录并写入签到时间；若已有记录但尚未签到，则补写签到时间；若当天已经签过到，则直接返回失败。同理，“下班签退”会补写或创建签退记录。",
    )
    add_text_paragraph(
        document,
        "这一实现方式简单直接，适合作为基础版本，但也说明当前系统的考勤规则仍较轻量，例如尚未看到固定班次、迟到早退自动计算、节假日规则或加班时长折算等复杂能力。",
    )
    add_heading(document, "7.2 请假申请与审批流程", 2)
    add_text_paragraph(
        document,
        "员工提交请假申请时，后端会自动将 `userId` 设置为当前登录用户，并写入创建时间，审批人、审批时间和审批意见会先置空，状态字段初始化为待审批。管理员后续可通过单条审批接口或批量审批接口修改状态，并记录审批意见与审批时间。",
    )
    add_heading(document, "7.3 薪资管理流程", 2)
    add_text_paragraph(
        document,
        "薪资模块当前由 `SalaryApiController` 直接使用 `JdbcTemplate` 操作 `emp_salary` 表。管理员可以根据用户名或用户 ID 录入某月薪资；员工则通过 `/api/salaries/mine` 查询属于自己的薪资记录。这种写法简洁，但从工程一致性上看，未来可以补齐 Salary 实体、Mapper 和 Service，使代码风格与其他模块统一。",
    )
    add_heading(document, "7.4 首页统计流程", 2)
    add_text_paragraph(
        document,
        "首页统计模块主要由 `/api/dashboard/summary`、`/dept-employee`、`/attendance-status` 和 `/add-leave-trend` 四类接口组成。前端 `WelcomePage.vue` 在页面挂载后同时拉取汇总指标和 ECharts 图表数据，并定时刷新。统计 SQL 直接落在 `DashboardApiController` 中，优点是开发快，缺点是后续一旦统计口径复杂，控制器容易膨胀。",
    )

    add_heading(document, "8. 前后端协作机制", 1)
    add_text_paragraph(
        document,
        "前端通过 Vite 开发服务器把 `/api` 请求代理到 `http://localhost:8080`，因此开发阶段 Vue 与 Spring Boot 可以分开启动而不需要手动处理跨域。与此同时，`application.yml` 与 `WebMvcConfig`、`SecurityConfig` 都显式配置了允许的本地来源地址，保证跨域请求可以携带 Cookie。",
    )
    for item in [
        "Axios 默认开启 `withCredentials: true`，确保浏览器自动发送 Session Cookie。",
        "前端统一使用 `unwrapResponse` 解析 `RetJson`，约定只有 `code = 200` 才视为成功。",
        "路由守卫会先执行 `restoreAuth()`，通过 `/api/auth/me` 恢复登录状态，然后再决定是否允许进入页面。",
        "管理员页面除了前端 `adminOnly` 限制外，后端也会再次校验，避免仅靠前端菜单隐藏造成越权。",
    ]:
        add_bullet(document, item)
    add_text_paragraph(
        document,
        "这种双重校验方式体现了一个比较正确的工程思路：前端负责优化体验，后端负责最终安全边界。即使有人绕过前端直接访问接口，后端依然会返回 401 或 403。",
    )

    add_heading(document, "9. 部署、运行与打包方式", 1)
    add_table(
        document,
        ["项目项", "要求或命令", "说明"],
        [
            ["JDK", "17+", "与 `pom.xml` 中的 `java.version` 保持一致"],
            ["Maven", "3.9+", "用于启动和打包 Spring Boot 项目"],
            ["Node.js", "20+", "用于前端依赖安装与 Vite 构建"],
            ["MySQL", "8+", "默认数据库名为 `ems-db`"],
            ["后端启动", "mvn spring-boot:run", "默认访问地址为 `http://localhost:8080`"],
            ["前端启动", "cd frontend-vue && npm install && npm run dev", "默认访问地址为 `http://localhost:5173/signin`"],
            ["后端打包", "mvn -DskipTests package", "生成 `target/attendance-management.war`"],
            ["前端打包", "cd frontend-vue && npm run build", "生成 `frontend-vue/dist`"],
        ],
        column_widths=[2.4, 6.2, 7.6],
    )
    add_text_paragraph(
        document,
        "数据库连接支持通过环境变量 `DB_URL`、`DB_USERNAME`、`DB_PASSWORD` 覆盖默认值，因此在本地、测试环境和答辩演示环境之间切换时，不必反复修改源码配置文件。这一点对项目演示和部署都比较友好。",
    )

    add_heading(document, "10. 当前实现亮点与优化建议", 1)
    add_heading(document, "10.1 亮点总结", 2)
    for item in [
        "项目已经完成从传统项目结构向 Spring Boot + Vue 前后端分离结构的迁移，整体方向正确。",
        "统一的 `RetJson` 和 `PageResult` 降低了前后端接口对接成本。",
        "Session 鉴权实现简单、成本低，适合管理后台型系统。",
        "密码校验支持明文与 BCrypt 共存，便于旧数据平滑迁移。",
        "管理员与员工使用同一套前端骨架，降低了页面重复开发量。",
    ]:
        add_bullet(document, item)
    add_heading(document, "10.2 优化建议", 2)
    add_table(
        document,
        ["方向", "当前现状", "建议做法"],
        [
            ["角色权限", "管理员身份仍依赖 `admin` 账号兼容判断", "引入用户、角色、权限、角色菜单映射，升级为标准 RBAC"],
            ["考勤规则", "当前仅记录签到签退时间，规则相对简单", "增加班次、迟到早退判定、节假日、补卡和加班规则"],
            ["数据访问一致性", "部分模块使用 MyBatis，部分模块直接写 JdbcTemplate", "逐步统一为实体 + Mapper + Service，减少 SQL 分散"],
            ["统计能力", "统计 SQL 直接放在 Controller 中", "沉淀为专门的报表服务，便于扩展月报、季报、导出功能"],
            ["测试与初始化", "未见自动化测试和初始化 SQL", "补充单元测试、接口测试和演示数据库脚本"],
            ["可审计性", "暂无操作日志、审批轨迹和导出记录", "增加审计日志、异常告警和数据导出记录"],
        ],
        column_widths=[2.6, 5.4, 8.2],
    )
    add_text_paragraph(
        document,
        "如果把这个系统作为课程设计项目，当前版本已经能够支撑完整答辩；如果把它作为真实业务系统继续演进，上述六个方向会是优先级较高的增强点。",
    )

    add_heading(document, "11. 关键源码入口参考", 1)
    add_table(
        document,
        ["关键位置", "职责说明"],
        [
            ["src/main/java/edu/com/eams/config/SecurityConfig.java", "定义接口访问规则、401/403 处理方式和 CORS 配置"],
            ["src/main/java/edu/com/eams/security/SessionAuthenticationFilter.java", "从 Session 恢复认证信息并注入 Spring Security 上下文"],
            ["src/main/java/edu/com/eams/controller/api/AuthApiController.java", "处理登录、退出和当前用户信息查询"],
            ["src/main/java/edu/com/eams/controller/api/AttendanceApiController.java", "处理考勤列表、详情、打卡和签退接口"],
            ["src/main/java/edu/com/eams/controller/api/LeaveApiController.java", "处理请假申请与审批相关接口"],
            ["src/main/java/edu/com/eams/controller/api/SalaryApiController.java", "处理薪资维护与个人薪资查询"],
            ["src/main/resources/mapper/*.xml", "定义 MyBatis SQL 和实体字段映射关系"],
            ["frontend-vue/src/router/index.js", "定义页面路由与访问守卫"],
            ["frontend-vue/src/stores/auth.js", "负责前端登录态恢复与角色判断"],
            ["frontend-vue/src/pages/WelcomePage.vue", "实现首页统计与图表展示"],
            ["frontend-vue/vite.config.js", "定义开发环境下前后端代理转发规则"],
        ],
        column_widths=[7.0, 9.2],
    )

    add_heading(document, "12. 结语", 1)
    add_text_paragraph(
        document,
        "综合来看，企业员工考察管理系统已经具备较完整的管理后台雏形：业务模块覆盖面较全，前后端职责划分较清楚，部署方式明确，权限控制也达到了可用水平。对于课程展示而言，这套系统已经能很好地体现 Java Web 项目的分层设计、数据库建模能力、前后端协作方式和实际业务流程理解。",
    )
    add_text_paragraph(
        document,
        "后续如果继续打磨，可以围绕“标准化权限、完整考勤规则、报表导出、测试体系、审批流程深化”继续增强。这样系统就能从一个可展示的教学项目，逐步演进为更接近真实企业内部应用的完整解决方案。",
    )

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
